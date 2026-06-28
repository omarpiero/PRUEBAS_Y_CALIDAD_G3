from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
SOURCE_DOCX = Path(r"C:\Users\Ani\Downloads\INFORME_FINAL.docx")
OUTPUT_DIR = REPO_ROOT / "documentacion" / "informe_final_actualizado"
OUTPUT_DOCX = OUTPUT_DIR / "INFORME_FINAL_ACTUALIZADO_G03.docx"


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_YELLOW = "FFF2CC"
LIGHT_RED = "FCE4D6"
GRAY = "F2F2F2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "A6A6A6")


def try_table_grid_style(table) -> None:
    try:
        table.style = "Table Grid"
    except KeyError:
        set_table_borders(table)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill: str = BLUE, body_fill: str | None = None) -> None:
    try_table_grid_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            elif body_fill:
                set_cell_shading(cell, body_fill)


def add_status_table(doc: Document) -> None:
    doc.add_heading("D.2 Estado de entregables frente a la consigna", level=2)
    rows = [
        (
            "Informe final completo",
            "Cubierto y ampliado",
            "Se conserva el informe base y se agrega este anexo de cierre para cubrir repositorio, despliegue, pruebas, evidencias, legal y reparto.",
        ),
        (
            "Enlace GIT y módulos",
            "Cubierto con pendiente de validación nominal",
            "Repositorio: https://github.com/omarpiero/PRUEBAS_Y_CALIDAD_G3.git. Falta que cada integrante confirme su módulo exacto si el docente lo exige por persona.",
        ),
        (
            "Aplicativo implementado",
            "Documentado",
            "El proyecto queda como LMS ejecutable localmente. No se identificó URL pública permanente; debe añadirse si el grupo despliega en hosting.",
        ),
        (
            "Diapositivas",
            "Cubierto",
            "Archivo: documentacion/presentacion_final/Presentacion_Final_LMS_G03.pptx.",
        ),
        (
            "Estado del arte",
            "Cubierto en informe base",
            "Mantener referencias usadas en el capítulo correspondiente y alinearlas con ISO 9001, ISO/IEC 25000, ISO/IEC/IEEE 29119 e ISO/IEC 27000.",
        ),
        (
            "Plan de pruebas",
            "Cubierto con restricción",
            "La implementación de pruebas queda documentada; la ejecución final depende de resolver Composer/vendor en el entorno local.",
        ),
        (
            "Diseño/mockups",
            "Documentado",
            "Se incorporó la ruta de cierre; falta exportar capturas finales de Figma si el docente exige evidencia visual dentro del informe.",
        ),
        (
            "Código y matriz de doble entrada",
            "Cubierto",
            "Se añadieron documentos de matriz, macroprocesos, procedimientos, actividades y funciones en documentacion/cierre_informe_final/.",
        ),
        (
            "Mantenimiento y SQ",
            "Cubierto",
            "Se documentó mantenimiento, seguimiento de calidad y cierre de DEF-001.",
        ),
        (
            "Despliegue 1ra versión",
            "Parcial",
            "Queda documentado como primera versión demostrable localmente; falta URL pública si el grupo decide publicar la aplicación.",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = ["Elemento solicitado", "Estado", "Detalle / evidencia"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for item, status, detail in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], item, bold=True)
        set_cell_text(cells[1], status)
        set_cell_text(cells[2], detail)
        if status.startswith("Cubierto"):
            set_cell_shading(cells[1], LIGHT_GREEN)
        elif status.startswith("Parcial"):
            set_cell_shading(cells[1], LIGHT_YELLOW)
        else:
            set_cell_shading(cells[1], GRAY)
    style_table(table)


def add_team_table(doc: Document) -> None:
    doc.add_heading("D.3 Reparto sugerido para sustentar y cerrar por integrantes", level=2)
    doc.add_paragraph(
        "La siguiente distribución permite que cada integrante defienda un módulo y una evidencia concreta. "
        "Los nombres pueden reemplazarse por los integrantes reales antes de la sustentación."
    )
    rows = [
        (
            "Integrante 1",
            "Administración de usuarios y roles",
            "Modelo User, controlador Admin/UserController, control de is_admin y prueba de asignación masiva.",
            "Explicar seguridad de roles, riesgo corregido DEF-001 y evidencia en Git.",
        ),
        (
            "Integrante 2",
            "Cursos / contenidos LMS",
            "Flujos de cursos, validaciones y vistas relacionadas con el aprendizaje.",
            "Mostrar caso de uso, pantalla, ruta de código y prueba asociada.",
        ),
        (
            "Integrante 3",
            "Evaluaciones / progreso",
            "Procesos de evaluación, seguimiento académico y persistencia de resultados.",
            "Defender matriz proceso-procedimiento-actividad-función.",
        ),
        (
            "Integrante 4",
            "Calidad, pruebas y documentación",
            "Plan de pruebas, matriz de doble entrada, estándares ISO y presentación final.",
            "Sustentar trazabilidad ISO 9001, 25000, 29119 y 27000.",
        ),
    ]
    table = doc.add_table(rows=1, cols=4)
    headers = ["Responsable", "Módulo", "Evidencia en proyecto", "Qué debe sustentar"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, bold=(i == 0))
    style_table(table)


def add_evidence_table(doc: Document) -> None:
    doc.add_heading("D.4 Evidencias generadas para el cierre", level=2)
    rows = [
        ("08_unificacion_metricas_pruebas.md", "Consolida métricas de pruebas y criterios de aceptación."),
        ("09_matriz_doble_entrada_cierre.md", "Relaciona macroprocesos, procedimientos, actividades, funciones y evidencias."),
        ("10_anexos_macroprocesos.md", "Amplía anexos operativos del proyecto."),
        ("11_cierre_defecto_def_001.md", "Documenta el defecto de asignación masiva de administrador y su corrección."),
        ("12_estado_legal_indecopi.md", "Detalla pendientes legales, titularidad, licencias y registro INDECOPI."),
        ("13_mantenimiento_sq_plan_pruebas.md", "Integra mantenimiento, SQ y plan de pruebas."),
        ("14_privacidad_terminos_licencias.md", "Define privacidad, términos y control de licencias."),
        ("15_conclusiones_recomendaciones_cierre.md", "Ajusta conclusiones y recomendaciones a un estado real de MVP."),
        ("16_repositorio_modulos_integrantes.md", "Ordena evidencia Git y módulos por integrante."),
        ("17_demo_url_despliegue.md", "Aclara estado de URL, demo local y despliegue."),
        ("18_mockups_figma_cierre.md", "Organiza evidencia de diseño y mockups."),
        ("19_guia_integracion_word_final.md", "Guía para integrar anexos al informe Word."),
        ("20_release_tag_zip.md", "Evidencia de release, tag y ZIP de entrega."),
        ("21_estado_pruebas_instalacion.md", "Registra bloqueo de instalación/pruebas por Composer/vendor."),
        ("Presentacion_Final_LMS_G03.pptx", "Diapositivas listas para sustentación."),
    ]
    table = doc.add_table(rows=1, cols=2)
    headers = ["Archivo", "Uso en informe final"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for file_name, usage in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], file_name, bold=True)
        set_cell_text(cells[1], usage)
    style_table(table)


def add_legal_section(doc: Document) -> None:
    doc.add_heading("D.5 Cierre legal y documental", level=2)
    rows = [
        ("Titularidad", "Confirmar por escrito quiénes son autores, responsables académicos y titulares del producto."),
        ("Cesión/permiso de uso", "Preparar autorización simple de uso académico o cesión si el producto se entrega a una entidad externa."),
        ("Licencias", "Mantener listado de dependencias Composer/NPM y declarar respeto de licencias de terceros."),
        ("Datos personales", "Documentar datos tratados por el LMS, finalidad, responsables y medidas básicas de seguridad."),
        ("INDECOPI", "Conservar autoría, repositorio, versión, capturas y ZIP si se busca registro de software."),
        ("Seguridad ISO/IEC 27000", "Preservar controles de roles, contraseñas, mínima exposición de datos y revisión de permisos."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_cell_text(table.rows[0].cells[0], "Aspecto", bold=True)
    set_cell_text(table.rows[0].cells[1], "Acción necesaria", bold=True)
    set_cell_text(table.rows[0].cells[2], "Estado", bold=True)
    for aspect, action in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], aspect, bold=True)
        set_cell_text(cells[1], action)
        set_cell_text(cells[2], "Por confirmar")
        set_cell_shading(cells[2], LIGHT_YELLOW)
    style_table(table)


def add_tests_section(doc: Document) -> None:
    doc.add_heading("D.6 Estado de pruebas y calidad", level=2)
    doc.add_paragraph(
        "El plan de pruebas queda documentado y trazado con ISO/IEC/IEEE 29119. "
        "También se cerró el hallazgo DEF-001 asociado a asignación masiva del rol administrador. "
        "La ejecución automatizada completa no pudo certificarse en este entorno porque Composer no terminó de instalar "
        "dependencias y no se generó vendor/autoload.php. Este bloqueo está registrado como evidencia técnica y debe "
        "resolverse antes de afirmar resultados finales de pruebas automatizadas."
    )
    rows = [
        ("DEF-001", "Control de is_admin", "Cerrado en código", "Modelo protegido y controlador con forceFill controlado."),
        ("Feature test", "Asignación masiva", "Añadido", "AdminSecurityAndRolesTest valida que is_admin no sea fillable."),
        ("Composer/vendor", "Entorno de pruebas", "Bloqueado", "Pendiente instalar dependencias sin error de permisos/red."),
        ("Pruebas finales", "Métricas", "Pendiente de ejecución", "Ejecutar php artisan test cuando vendor/autoload.php exista."),
    ]
    table = doc.add_table(rows=1, cols=4)
    headers = ["Elemento", "Alcance", "Estado", "Observación"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, bold=(i == 0))
        if row[2] == "Bloqueado":
            set_cell_shading(cells[2], LIGHT_RED)
        elif row[2].startswith("Pendiente"):
            set_cell_shading(cells[2], LIGHT_YELLOW)
        else:
            set_cell_shading(cells[2], LIGHT_GREEN)
    style_table(table)


def add_deploy_section(doc: Document) -> None:
    doc.add_heading("D.7 Despliegue de primera versión", level=2)
    doc.add_paragraph(
        "La primera versión queda evidenciada como producto de software integrado en Git y preparado para demostración local. "
        "Si el grupo publica el aplicativo en Render, Railway, Vercel, hosting institucional u otro servicio, debe reemplazar "
        "este texto por la URL pública y anexar una captura de disponibilidad."
    )
    rows = [
        ("Repositorio", "https://github.com/omarpiero/PRUEBAS_Y_CALIDAD_G3.git"),
        ("Rama de cierre", "main"),
        ("Tag de cierre creado", "informe-final-g03-2026-06-27"),
        ("Presentación", "documentacion/presentacion_final/Presentacion_Final_LMS_G03.pptx"),
        ("Demo pública", "No identificada al cierre; pendiente si el docente exige URL web."),
        ("Demo local", "Configurar .env, base de datos, composer install, migraciones y php artisan serve."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.rows[0].cells[0], "Campo", bold=True)
    set_cell_text(table.rows[0].cells[1], "Valor")
    for field, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], field, bold=True)
        set_cell_text(cells[1], value)
    style_table(table)


def add_closing_checklist(doc: Document) -> None:
    doc.add_heading("D.8 Checklist final para el jefe de grupo", level=2)
    checks = [
        "Confirmar nombres reales y módulos defendidos por cada integrante.",
        "Pegar capturas finales de mockups/Figma en el informe si todavía no aparecen.",
        "Agregar URL pública del aplicativo si se despliega antes de la sustentación.",
        "Resolver Composer/vendor y ejecutar php artisan test para reemplazar el estado bloqueado por resultados reales.",
        "Actualizar índice de Word si el documento mantiene tabla de contenido automática.",
        "Verificar que la presentación y el informe tengan los mismos nombres de proyecto, integrantes y evidencias.",
        "Conservar ZIP de entrega, tag Git y hash SHA-256 como evidencia de versión.",
    ]
    table = doc.add_table(rows=1, cols=3)
    set_cell_text(table.rows[0].cells[0], "Nro.", bold=True)
    set_cell_text(table.rows[0].cells[1], "Acción", bold=True)
    set_cell_text(table.rows[0].cells[2], "Estado sugerido", bold=True)
    for index, check in enumerate(checks, start=1):
        cells = table.add_row().cells
        set_cell_text(cells[0], str(index), bold=True)
        set_cell_text(cells[1], check)
        set_cell_text(cells[2], "Pendiente / validar")
        set_cell_shading(cells[2], LIGHT_YELLOW)
    style_table(table)


def normalize_new_section(doc: Document) -> None:
    section = doc.sections[-1]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def build() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"No se encontro el informe base: {SOURCE_DOCX}")
    shutil.copy2(SOURCE_DOCX, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)

    doc.add_section(WD_SECTION.NEW_PAGE)
    normalize_new_section(doc)

    title = doc.add_heading("Anexo D: Cierre de pendientes finales del proyecto", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor.from_string(BLUE)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead.add_run(
        "Este anexo integra los elementos que faltaban para que el informe final quede completo frente a la consigna de la asignatura. "
    ).bold = True
    lead.add_run(
        "No reemplaza el desarrollo previo: lo complementa con evidencias de Git, despliegue, pruebas, documentación legal, "
        "presentación y tareas de cierre que el grupo debe sustentar."
    )

    doc.add_heading("D.1 Resumen ejecutivo de cierre", level=2)
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary.add_run("Proyecto: ").bold = True
    summary.add_run("Sistema LMS JMJS del Grupo 03. ")
    summary.add_run("Enfoque de calidad: ").bold = True
    summary.add_run(
        "ISO 9001 para gestión del proceso, ISO/IEC 25000 para calidad del producto, "
        "ISO/IEC/IEEE 29119 para pruebas e ISO/IEC 27000 para seguridad de la información."
    )

    add_status_table(doc)
    add_team_table(doc)
    add_evidence_table(doc)
    add_legal_section(doc)
    add_tests_section(doc)
    add_deploy_section(doc)
    add_closing_checklist(doc)

    doc.add_paragraph(
        "Nota de control documental: este anexo fue generado para la versión de cierre del informe final y debe revisarse "
        "junto con las evidencias versionadas en Git antes de la sustentación."
    )

    doc.core_properties.title = "Informe final actualizado - LMS JMJS G03"
    doc.core_properties.subject = "Cierre de pendientes de documentacion, pruebas, despliegue y legal"
    doc.core_properties.keywords = "ISO 9001, ISO 25000, ISO 29119, ISO 27000, LMS, pruebas, calidad"
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build()
